import json
import re
import pandas as pd
from groq import Groq
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import os
from ddgs import DDGS

client = Groq(api_key="dont see brouououou")

instrucciones_excel = """
Respond ONLY with valid JSON.

Structure:
{
  "datos": [[...], [...]],
  "columnas": ["col1", "col2"],
  "estilo": {
    "font_size": <number, default 11>,
    "header_color": <hex sin #, default "2F4F7F">,
    "font_color_header": <hex sin #, default "FFFFFF">,
    "row_alt_color": <hex sin #, default "F2F2F2">
  }
}

Rules:
- Only include "grafico" field if the user explicitly asks for a chart
- "grafico" types: "bar", "line", "pie"
- "columna_x" and "columna_y" must match exact column names in "columnas"
- If a calculated column is needed, use EXACTLY these names: "Total", "Subtotal", "Promedio", "Cantidad"
- "Total" = price × quantity
- "Subtotal" = same as Total before taxes
- "Promedio" = average of a numeric column
- You MUST use the real data provided under "Real data found" as the primary source. Never invent information that contradicts it.
- don't make up or say anything that isn't proven
- No explanations, no markdown, no extra text
- columnas y datos must match in quantity
- Descriptions must be detailed, specific and distinct from each other
- Data must be realistic and coherent with the context
- Never repeat the same value in a description column
- Style must be coherent with the requested Excel theme
- All text content (column names, data) must be written in Spanish
- NEVER fill price or monetary columns with values, always set them to 0. This is mandatory.
"""

instrucciones_word = """
Respond ONLY with valid JSON.

Structure:
{
  "titulo": "Document title in Spanish",
  "terminos": [
    {
      "nombre": "Concept name in Spanish",
      "definicion": "Detailed definition in Spanish.",
      "palabras_clave": ["word1", "word2"]
    }
  ]
}

Rules:
- You MUST use the real data provided under "Real data found" as the primary source. Never invent information that contradicts it.
- No explanations, no markdown, no extra text
- Definitions must be clear, complete and academic, written in Spanish
- palabras_clave are the most important terms within the definition (2-4 per concept)
- Never repeat definitions
- don't make up or say anything that isn't proven
"""

def iu_basica():
    seleccion = input("Introduzca 1 para generar un Excel y 2 para un Word: ")
    if seleccion == "1":
        prompt = input("Describa el Excel: ")
    else:
        prompt = input("Describa el Word: ")
    return prompt, seleccion


def mejorar_prompt(prompt_usuario, tipo):
    tipo_str = "Excel con datos tabulares" if tipo == "1" else "Word con definiciones"

    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": f"""
You are an assistant that improves prompts for generating {tipo_str}.
Take the user input and rewrite it clearly, specifically and in detail.
Add missing context, specify number of columns/terms if not mentioned.
Always request detailed descriptions regardless of what the user says.
The final output (Excel/Word content) must be in Spanish.
Return ONLY the improved prompt, no explanations or comments.
"""},
            {"role": "user", "content": prompt_usuario}
        ],
        temperature=0.3
    )
    return completion.choices[0].message.content

def obtener_precios_meli(prompt_original):
    # Usamos el prompt original del usuario, no el mejorado
    # Es corto y tiene exactamente lo que necesitamos
    palabras = [p.strip(",.") for p in prompt_original.split() if len(p) > 3]
    
    resultados = []
    for termino in palabras:
        try:
            url = f"https://api.mercadolibre.com/sites/MLA/search?q=yerba+{termino}&limit=1"
            r = requests.get(url, timeout=5)
            data = r.json()
            if data["results"]:
                item = data["results"][0]
                resultados.append(f"{item['title']}: ${item['price']}")
        except Exception:
            pass

    return "\n".join(resultados) if resultados else "No se encontraron precios."

def generacion_json(prompt, instrucciones):
    completion = client.chat.completions.create(
        model="openai/gpt-oss-120b", # otra opcion valida es llama-3.3-70b-versatile otro es openai/gpt-oss-120b otra es groq/compound 
        messages=[
            {"role": "system", "content": instrucciones},
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )
    return completion.choices[0].message.content

def parsear_json(texto):
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        pass
    match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", texto, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    match = re.search(r"\{.*\}", texto, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass
    return None




def buscar_datos_web(prompt_original):
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": "Extract ONLY the product or brand names from the text. Return them separated by commas, nothing else. Example output: taragui, rosamonte, cbse"},
            {"role": "user", "content": prompt_original}
        ],
        temperature=0,
        max_tokens=30
    )
    keywords = completion.choices[0].message.content.strip()
    print("KEYWORDS:", keywords)

    resultados = []
    for kw in [k.strip() for k in keywords.split(",")]:
        results = DDGS().text(f"{kw} yerba mate Argentina", max_results=2)
        for r in results:
            resultados.append(r["body"])

    return "\n".join(resultados[:6]) if resultados else "No se encontraron datos."

def formatear_excel(path, estilo):
    wb = load_workbook(path)
    ws = wb.active

    header_fill = PatternFill("solid", fgColor=estilo.get("header_color", "2F4F7F"))
    header_font = Font(bold=True, color=estilo.get("font_color_header", "FFFFFF"), size=estilo.get("font_size", 11))

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    fill_alt = PatternFill("solid", fgColor=estilo.get("row_alt_color", "F2F2F2"))
    for i, row in enumerate(ws.iter_rows(min_row=2), start=2):
        for cell in row:
            if i % 2 == 0:
                cell.fill = fill_alt
            cell.font = Font(size=estilo.get("font_size", 11))
            cell.alignment = Alignment(horizontal="left", wrap_text=True)

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.border = border

    aplicar_formulas(ws)
    wb.save(path)
    wb.save(path)

def aplicar_formulas(ws):
    headers = [cell.value for cell in ws[1]]

    def buscar_col(keywords):
        for i, h in enumerate(headers):
            if h and any(k.lower() in h.lower() for k in keywords):
                return i + 1
        return None

    col_total = buscar_col(["total", "subtotal"])
    col_precio = buscar_col(["precio", "price", "unitario"])
    col_cantidad = buscar_col(["cantidad", "quantity", "vendida"])
    col_promedio = buscar_col(["promedio", "average"])

    ultima_fila = ws.max_row

    # Fórmula Total = Precio * Cantidad por fila
    for row in range(2, ultima_fila + 1):
        if col_total and col_precio and col_cantidad:
            precio_cell = ws.cell(row=row, column=col_precio).coordinate
            cantidad_cell = ws.cell(row=row, column=col_cantidad).coordinate
            ws.cell(row=row, column=col_total).value = f"={precio_cell}*{cantidad_cell}"

    # Fila de totales al final
    fila_suma = ultima_fila + 2  # deja una fila vacía de separación

    if col_total:
        col_total_letra = ws.cell(row=2, column=col_total).column_letter
        ws.cell(row=fila_suma, column=col_total).value = f"=SUM({col_total_letra}2:{col_total_letra}{ultima_fila})"
        ws.cell(row=fila_suma, column=col_total - 1).value = "TOTAL"

    if col_cantidad:
        col_cantidad_letra = ws.cell(row=2, column=col_cantidad).column_letter
        ws.cell(row=fila_suma, column=col_cantidad).value = f"=SUM({col_cantidad_letra}2:{col_cantidad_letra}{ultima_fila})"

    if col_promedio:
        col_promedio_letra = ws.cell(row=2, column=col_promedio).column_letter
        ws.cell(row=fila_suma, column=col_promedio).value = f"=AVERAGE({col_promedio_letra}2:{col_promedio_letra}{ultima_fila})"
        ws.cell(row=fila_suma, column=col_promedio - 1).value = "PROMEDIO"

    # Estilo de la fila de totales
    from openpyxl.styles import Font
    for col in range(1, ws.max_column + 1):
        cell = ws.cell(row=fila_suma, column=col)
        cell.font = Font(bold=True, size=11)

def generar_word(data, output_path):
    titulo = data.get("titulo", "Documento")
    terminos = data.get("terminos", [])

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    titulo_par = doc.add_paragraph()
    titulo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_par.paragraph_format.space_after = Pt(24)
    run_titulo = titulo_par.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(24)
    run_titulo.font.name = "Arial"

    for t in terminos:
        nombre = t.get("nombre", "")
        definicion = t.get("definicion", "")
        palabras_clave = [pk.lower() for pk in t.get("palabras_clave", [])]

        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(10)
        par.paragraph_format.space_after = Pt(10)

        run_nombre = par.add_run(nombre + ": ")
        run_nombre.bold = True
        run_nombre.font.size = Pt(12)
        run_nombre.font.name = "Arial"

        restante = definicion
        while restante:
            primer_idx = len(restante)
            primer_palabra = None

            for pk in palabras_clave:
                idx = restante.lower().find(pk)
                if idx != -1 and idx < primer_idx:
                    primer_idx = idx
                    primer_palabra = pk

            if primer_palabra is None:
                run = par.add_run(restante)
                run.font.size = Pt(12)
                run.font.name = "Arial"
                break
            else:
                if primer_idx > 0:
                    run = par.add_run(restante[:primer_idx])
                    run.font.size = Pt(12)
                    run.font.name = "Arial"

                run_sub = par.add_run(restante[primer_idx:primer_idx + len(primer_palabra)])
                run_sub.underline = True
                run_sub.font.size = Pt(12)
                run_sub.font.name = "Arial"

                restante = restante[primer_idx + len(primer_palabra):]

    doc.save(output_path)
    print("Word creado")


# MAIN
prompt, seleccion = iu_basica()

prompt_mejorado = mejorar_prompt(prompt, seleccion)
print("PROMPT MEJORADO:", prompt_mejorado)
print("=" * 50)

print("Buscando datos en la web...")
datos_web = buscar_datos_web(prompt)
print("DATOS WEB:", datos_web)
print("=" * 50)

if seleccion == "1":
    contenido = generacion_json(f"{prompt_mejorado}\n\nReal data found:\n{datos_web}", instrucciones_excel)
else:
    contenido = generacion_json(f"{prompt_mejorado}\n\nReal data found:\n{datos_web}", instrucciones_word)

print("RESPUESTA CRUDA:")
print(contenido)
print("=" * 50)

data = parsear_json(contenido)
if data is None:
    print("No se pudo extraer JSON válido")
    input("Presioná Enter para cerrar...")
    exit()

if seleccion == "1":
    try:
        nombre = input("¿Qué nombre le ponés al archivo Excel? (sin extensión): ")
        os.makedirs("excels", exist_ok=True)
        path = f"excels/{nombre}.xlsx"
        df = pd.DataFrame(data["datos"], columns=data["columnas"])
        df.to_excel(path, index=False)
        estilo = data.get("estilo", {})
        formatear_excel(path, estilo)
        print(f"Excel creado: {path}")
        print("NOTA, MUY IMPORTANTE POR FAVOR LEEEEEEEEEEEEEEEEEEEEEEEEEEEEEER: holaaa soy fran porfa porfa revisa lo que haya generado  porque se puede equivocar, el modulo con el modelo de ia que busca la info esta medio gaga aveces, gracias x usar:D ")
    except Exception as e:
        print("Error:", e)
else:
    try:
        nombre = input("¿Qué nombre le ponés al archivo Word? (sin extensión): ")
        os.makedirs("words", exist_ok=True)
        path = f"words/{nombre}.docx"
        generar_word(data, path)
    except Exception as e:
        print("Error:", e)