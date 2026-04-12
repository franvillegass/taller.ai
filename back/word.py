from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_word(data, output_path):
    titulo = data.get("titulo", "Documento")
    terminos = data.get("terminos", [])

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    titulo_par = doc.add_paragraph()
    titulo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_par.paragraph_format.space_after = Pt(24)
    run_titulo = titulo_par.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(24)
    run_titulo.font.name = "Arial"

    for t in terminos:
        nombre = t.get("nombre", "")
        definicion = t.get("definicion", "")
        palabras_clave = [pk.lower() for pk in t.get("palabras_clave", [])]

        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(10)
        par.paragraph_format.space_after = Pt(10)

        run_nombre = par.add_run(nombre + ": ")
        run_nombre.bold = True
        run_nombre.font.size = Pt(12)
        run_nombre.font.name = "Arial"

        restante = definicion
        while restante:
            primer_idx = len(restante)
            primer_palabra = None

            for pk in palabras_clave:
                idx = restante.lower().find(pk)
                if idx != -1 and idx < primer_idx:
                    primer_idx = idx
                    primer_palabra = pk

            if primer_palabra is None:
                run = par.add_run(restante)
                run.font.size = Pt(12)
                run.font.name = "Arial"
                break
            else:
                if primer_idx > 0:
                    run = par.add_run(restante[:primer_idx])
                    run.font.size = Pt(12)
                    run.font.name = "Arial"

                run_sub = par.add_run(restante[primer_idx:primer_idx + len(primer_palabra)])
                run_sub.underline = True
                run_sub.font.size = Pt(12)
                run_sub.font.name = "Arial"

                restante = restante[primer_idx + len(primer_palabra):]

    doc.save(output_path)
    print("Word creado")